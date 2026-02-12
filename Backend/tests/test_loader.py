"""
Unit Tests for Document Loader (Factory Pattern).

Tests cover:
    - Factory correctly selects the right loader.
    - TxtLoader reads plain text files.
    - Unsupported formats raise ValueError.
    - Missing files raise FileNotFoundError.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from app.services.document_loader import (
    BaseDocumentLoader,
    DocumentLoaderFactory,
    PDFLoader,
    DocxLoader,
    TxtLoader,
)


class TestDocumentLoaderFactory:
    """Tests for the Factory Pattern implementation."""

    def test_factory_returns_pdf_loader(self, tmp_path: Path) -> None:
        """Factory should return PDFLoader for .pdf files."""
        # Create a dummy file (won't be read in this test)
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4 dummy")

        loader = DocumentLoaderFactory.create_loader(str(pdf_file))
        assert isinstance(loader, PDFLoader)

    def test_factory_returns_docx_loader(self, tmp_path: Path) -> None:
        """Factory should return DocxLoader for .docx files."""
        # Create a minimal .docx file (a valid zip with [Content_Types].xml)
        from docx import Document

        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(docx_file))

        loader = DocumentLoaderFactory.create_loader(str(docx_file))
        assert isinstance(loader, DocxLoader)

    def test_factory_returns_txt_loader(self, tmp_path: Path) -> None:
        """Factory should return TxtLoader for .txt files."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello, world!", encoding="utf-8")

        loader = DocumentLoaderFactory.create_loader(str(txt_file))
        assert isinstance(loader, TxtLoader)

    def test_factory_raises_for_unsupported_format(self, tmp_path: Path) -> None:
        """Factory should raise ValueError for unsupported extensions."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("a,b,c", encoding="utf-8")

        with pytest.raises(ValueError, match="Unsupported file format"):
            DocumentLoaderFactory.create_loader(str(csv_file))

    def test_factory_raises_for_missing_file(self) -> None:
        """Factory should raise FileNotFoundError if file doesn't exist."""
        with pytest.raises(FileNotFoundError):
            DocumentLoaderFactory.create_loader("/nonexistent/resume.pdf")

    def test_supported_formats(self) -> None:
        """supported_formats() should return known extensions."""
        formats = DocumentLoaderFactory.supported_formats()
        assert ".pdf" in formats
        assert ".docx" in formats
        assert ".txt" in formats


class TestTxtLoader:
    """Tests for the TxtLoader concrete class."""

    def test_load_reads_content(self, tmp_path: Path) -> None:
        """TxtLoader should read the full content of a .txt file."""
        txt_file = tmp_path / "sample.txt"
        content = "John Doe\nSoftware Engineer\n5 years experience in Python"
        txt_file.write_text(content, encoding="utf-8")

        loader = TxtLoader(txt_file)
        result = loader.load()

        assert result == content
        assert "John Doe" in result
        assert "Python" in result

    def test_load_strips_whitespace(self, tmp_path: Path) -> None:
        """TxtLoader should strip leading/trailing whitespace."""
        txt_file = tmp_path / "padded.txt"
        txt_file.write_text("  \n  Hello World  \n  ", encoding="utf-8")

        loader = TxtLoader(txt_file)
        result = loader.load()

        assert result == "Hello World"


class TestDocxLoader:
    """Tests for the DocxLoader concrete class."""

    def test_load_reads_paragraphs(self, tmp_path: Path) -> None:
        """DocxLoader should extract text from all paragraphs."""
        from docx import Document

        docx_file = tmp_path / "resume.docx"
        doc = Document()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("Data Scientist")
        doc.add_paragraph("Machine Learning, Python, SQL")
        doc.save(str(docx_file))

        loader = DocxLoader(docx_file)
        result = loader.load()

        assert "Jane Smith" in result
        assert "Data Scientist" in result
        assert "Python" in result
